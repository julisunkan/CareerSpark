import os
import PyPDF2
import docx

def extract_text_from_file(filepath: str) -> str:
    """Extract text from PDF, DOCX, or TXT files"""
    
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    file_extension = filepath.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            return extract_text_from_pdf(filepath)
        elif file_extension == 'docx':
            return extract_text_from_docx(filepath)
        elif file_extension == 'txt':
            return extract_text_from_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        raise Exception(f"Error extracting text from {file_extension.upper()} file: {str(e)}")

def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from PDF file using PyPDF2"""
    text = ""
    
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            raise Exception("No readable text found in PDF. The PDF might be image-based or corrupted.")
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(filepath: str) -> str:
    """Extract text from DOCX file using python-docx"""
    try:
        doc = docx.Document(filepath)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise Exception("No readable text found in DOCX file.")
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_text_from_txt(filepath: str) -> str:
    """Extract text from TXT file"""
    try:
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text.strip():
                        return text.strip()
                    else:
                        raise Exception("Empty text file.")
            except UnicodeDecodeError:
                continue
        
        raise Exception("Could not decode text file with any common encoding.")
    
    except Exception as e:
        raise Exception(f"TXT extraction failed: {str(e)}")

def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    import re
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Remove common PDF artifacts
    text = text.replace('\uf0b7', '•')  # Replace bullet point artifacts
    text = text.replace('\x00', '')    # Remove null characters
    
    return text.strip()
